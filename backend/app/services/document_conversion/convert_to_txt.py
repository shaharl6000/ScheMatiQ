#!/usr/bin/env python3
"""
Multi-format to TXT Converter

Converts RTF, PDF, DOCX, and DOC files to plain text format.
- RTF/DOCX/DOC: Uses LibreOffice headless conversion
- PDF: Uses pdfplumber for text extraction

Usage:
    python -m src.preprocessing.convert_to_txt <input_dir> --output-dir <output_dir>

Example:
    python -m src.preprocessing.convert_to_txt "data/Cases to RK" --output-dir "data/US_Immigration/USA/raw"
"""

import argparse
import atexit
import os
import platform
import shutil
import signal
import subprocess
import sys
import tempfile
import time
import uuid
from concurrent.futures import ProcessPoolExecutor, as_completed
from pathlib import Path
from typing import Tuple

# Supported extensions
LIBREOFFICE_EXTENSIONS = {".rtf", ".docx", ".doc"}
PDF_EXTENSION = ".pdf"
SUPPORTED_EXTENSIONS = LIBREOFFICE_EXTENSIONS | {PDF_EXTENSION}

# Global temp directory for LibreOffice profiles
_temp_profile_base = None


def get_temp_profile_dir() -> Path:
    """
    Get or create the base temp directory for LibreOffice profiles.
    """
    global _temp_profile_base
    if _temp_profile_base is None:
        _temp_profile_base = Path(tempfile.gettempdir()) / f"libreoffice_profiles_{os.getpid()}"
        _temp_profile_base.mkdir(parents=True, exist_ok=True)
    return _temp_profile_base


def cleanup_temp_profiles():
    """
    Clean up temporary LibreOffice profile directories.
    """
    global _temp_profile_base
    if _temp_profile_base and _temp_profile_base.exists():
        try:
            shutil.rmtree(_temp_profile_base, ignore_errors=True)
        except Exception:
            pass


def kill_libreoffice_processes():
    """
    Kill all LibreOffice processes to ensure cleanup.
    Works on macOS, Linux, and Windows.
    """
    system = platform.system().lower()
    
    try:
        if system == "darwin":  # macOS
            # Kill soffice and soffice.bin processes
            subprocess.run(
                ["pkill", "-f", "soffice"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=5,
            )
        elif system == "linux":
            # Kill soffice and soffice.bin processes
            subprocess.run(
                ["pkill", "-f", "soffice"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=5,
            )
        elif system == "windows":
            # On Windows, use taskkill
            subprocess.run(
                ["taskkill", "/F", "/IM", "soffice.exe", "/T"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=5,
            )
            subprocess.run(
                ["taskkill", "/F", "/IM", "soffice.bin.exe", "/T"],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=5,
            )
    except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
        # Silently ignore errors - process might not exist
        pass


def signal_handler(signum, frame):
    """
    Handle interrupt signals (Ctrl+C) to ensure cleanup.
    """
    print("\n\nInterrupted! Cleaning up...")
    kill_libreoffice_processes()
    cleanup_temp_profiles()
    sys.exit(130)  # Standard exit code for SIGINT


def get_libreoffice_path() -> str:
    """
    Get the LibreOffice executable path based on the operating system.

    Returns:
        str: Path to the LibreOffice executable

    Raises:
        FileNotFoundError: If LibreOffice is not found
    """
    paths = {
        "darwin": "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "linux": "/usr/bin/soffice",
        "win32": "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
    }

    platform = sys.platform
    default_path = paths.get(platform, "/usr/bin/soffice")

    if os.path.exists(default_path):
        return default_path

    soffice = shutil.which("soffice")
    if soffice:
        return soffice

    raise FileNotFoundError(
        "LibreOffice not found. Please install LibreOffice or provide the path manually."
    )


def convert_with_libreoffice(input_path: Path, output_dir: Path, soffice_path: str, worker_id: str = None, max_retries: int = 3) -> Tuple[bool, str]:
    """
    Convert a file to TXT using LibreOffice headless with retry logic.

    Args:
        input_path: Path to the input file (RTF, DOCX, DOC)
        output_dir: Directory for output
        soffice_path: Path to LibreOffice executable
        worker_id: Unique worker ID for parallel processing (uses separate profile)
        max_retries: Maximum number of retry attempts (default: 3)

    Returns:
        Tuple of (success, message)
    """
    # Create a unique user profile directory for this worker to avoid conflicts
    if worker_id:
        profile_dir = get_temp_profile_dir() / f"profile_{worker_id}"
        profile_dir.mkdir(parents=True, exist_ok=True)
        user_installation = f"-env:UserInstallation=file://{profile_dir}"
    else:
        user_installation = None

    command = [
        soffice_path,
        "--headless",
    ]
    
    if user_installation:
        command.append(user_installation)
    
    command.extend([
        "--convert-to",
        "txt:Text",
        "--outdir",
        str(output_dir),
        str(input_path),
    ])

    last_error = ""
    for attempt in range(max_retries):
        try:
            result = subprocess.run(
                command,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
            )

            expected_output = output_dir / (input_path.stem + ".txt")

            if result.returncode == 0 and expected_output.exists():
                return True, f"Converted: {expected_output}"
            else:
                last_error = result.stderr.decode() if result.stderr else "Unknown error"
                
                # If this is not the last attempt, wait and retry
                if attempt < max_retries - 1:
                    time.sleep(2 * (attempt + 1))  # Exponential backoff: 2s, 4s
                    continue

        except subprocess.TimeoutExpired:
            last_error = "Conversion timed out after 120 seconds"
            if attempt < max_retries - 1:
                time.sleep(2 * (attempt + 1))
                continue
        except Exception as e:
            last_error = str(e)
            if attempt < max_retries - 1:
                time.sleep(2 * (attempt + 1))
                continue

    return False, f"LibreOffice conversion failed after {max_retries} attempts: {last_error}"


def convert_pdf_to_txt(input_path: Path, output_dir: Path) -> Tuple[bool, str]:
    """
    Convert a PDF to TXT. Tries pdfplumber first; falls back to OCR for scanned PDFs.
    Scanned pages yield <200 chars/page on average, which triggers the OCR path.
    """
    import pdfplumber
    from pdf2image import convert_from_path
    import pytesseract

    output_path = output_dir / (input_path.stem + ".txt")

    with pdfplumber.open(input_path) as pdf:
        page_count = max(1, len(pdf.pages))
        text_parts = [p.extract_text() or "" for p in pdf.pages]

    total_chars = sum(len(t) for t in text_parts)

    if total_chars / page_count >= 200:
        output_path.write_text("\n\n".join(t for t in text_parts if t), encoding="utf-8")
        return True, f"Converted: {output_path}"

    # Scanned PDF — use OCR
    pages = convert_from_path(str(input_path), dpi=300)
    ocr_parts = [pytesseract.image_to_string(img) for img in pages]
    ocr_parts = [t for t in ocr_parts if t.strip()]

    if not ocr_parts:
        return False, "No text extracted from PDF (tried pdfplumber and OCR)"

    output_path.write_text("\n\n".join(ocr_parts), encoding="utf-8")
    return True, f"Converted (OCR): {output_path}"


def convert_docx_to_txt(input_path: Path, output_dir: Path) -> Tuple[bool, str]:
    """
    Convert a DOCX file to TXT using python-docx.

    Args:
        input_path: Path to the DOCX file
        output_dir: Directory for output

    Returns:
        Tuple of (success, message)
    """
    try:
        from docx import Document
    except ImportError:
        return False, "python-docx not installed. Run: pip install python-docx"

    output_path = output_dir / (input_path.stem + ".txt")

    try:
        document = Document(input_path)
        text_parts = [para.text for para in document.paragraphs]

        # Include table text as well
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)

        if not text_parts:
            return False, "No text extracted from DOCX"

        output_path.write_text("\n".join(text_parts), encoding="utf-8")
        return True, f"Converted: {output_path}"

    except Exception as e:
        return False, f"DOCX extraction failed: {e}"


def convert_file(input_path: Path, output_dir: Path, soffice_path: str, worker_id: str = None) -> Tuple[bool, str]:
    """
    Convert a single file to TXT based on its extension.

    Args:
        input_path: Path to the input file
        output_dir: Directory for output
        soffice_path: Path to LibreOffice executable
        worker_id: Unique worker ID for parallel processing

    Returns:
        Tuple of (success, message)
    """
    ext = input_path.suffix.lower()

    if ext == ".docx":
        success, message = convert_docx_to_txt(input_path, output_dir)
        if success:
            return success, message
        # Fallback to LibreOffice if python-docx failed
        return convert_with_libreoffice(input_path, output_dir, soffice_path, worker_id)
    elif ext in {".doc", ".rtf"}:
        return convert_with_libreoffice(input_path, output_dir, soffice_path, worker_id)
    elif ext == PDF_EXTENSION:
        return convert_pdf_to_txt(input_path, output_dir)
    else:
        return False, f"Unsupported file type: {ext}"


def process_single_file(args_tuple: Tuple[Path, Path, str, bool]) -> Tuple[str, bool, str]:
    """
    Worker function for parallel processing.
    
    Args:
        args_tuple: Tuple of (input_path, output_dir, soffice_path, skip_existing)
    
    Returns:
        Tuple of (file_name, success, message)
    """
    input_path, output_dir, soffice_path, skip_existing = args_tuple
    
    output_path = output_dir / (input_path.stem + ".txt")
    
    if skip_existing and output_path.exists():
        return (input_path.name, True, "Skipped (exists)")
    
    # Generate unique worker ID for this process to avoid LibreOffice profile conflicts
    worker_id = f"{os.getpid()}_{uuid.uuid4().hex[:8]}"
    
    success, message = convert_file(input_path, output_dir, soffice_path, worker_id)
    return (input_path.name, success, message)


def process_directory(
    input_dir: Path,
    output_dir: Path,
    soffice_path: str,
    skip_existing: bool = True,
    num_workers: int = 16,
    libreoffice_workers: int = 2,
) -> dict:
    """
    Process all supported files in a directory using parallel processing.
    
    PDFs are processed with full parallelism.
    LibreOffice files (RTF/DOCX/DOC) are processed with limited workers to avoid conflicts.

    Args:
        input_dir: Directory containing source files
        output_dir: Directory for output TXT files
        soffice_path: Path to LibreOffice executable
        skip_existing: Skip files that already have a TXT output
        num_workers: Number of parallel workers for PDF files (default: 16)
        libreoffice_workers: Number of parallel workers for LibreOffice files (default: 2)

    Returns:
        Statistics dictionary
    """
    stats = {"total": 0, "successful": 0, "failed": 0, "skipped": 0}

    if not input_dir.exists():
        print(f"Input directory not found: {input_dir}")
        return stats

    output_dir.mkdir(parents=True, exist_ok=True)

    all_files = [
        f
        for f in input_dir.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS and not f.name.startswith("~$")
    ]

    if not all_files:
        print(f"No supported files found in {input_dir}")
        return stats

    # Separate PDF files from LibreOffice files
    pdf_files = [f for f in all_files if f.suffix.lower() == PDF_EXTENSION]
    libreoffice_files = [f for f in all_files if f.suffix.lower() in LIBREOFFICE_EXTENSIONS]

    stats["total"] = len(all_files)
    print(f"Found {stats['total']} file(s) to convert:")
    print(f"  - {len(pdf_files)} PDF files (using {num_workers} workers)")
    print(f"  - {len(libreoffice_files)} LibreOffice files (using {libreoffice_workers} workers)")

    def process_batch(files: list, workers: int, file_type: str):
        """Process a batch of files with specified number of workers."""
        if not files:
            return
        
        print(f"\nProcessing {len(files)} {file_type} files...")
        
        tasks = [
            (file_path, output_dir, soffice_path, skip_existing)
            for file_path in sorted(files)
        ]

        with ProcessPoolExecutor(max_workers=workers) as executor:
            future_to_file = {
                executor.submit(process_single_file, task): task[0]
                for task in tasks
            }

            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    file_name, success, message = future.result()
                    
                    if message == "Skipped (exists)":
                        print(f"  Skipping (exists): {file_name}")
                        stats["skipped"] += 1
                    elif success:
                        print(f"  OK: {file_name}")
                        stats["successful"] += 1
                    else:
                        print(f"  FAIL: {file_name} - {message}")
                        stats["failed"] += 1
                except Exception as e:
                    print(f"  ERROR: {file_path.name} - {e}")
                    stats["failed"] += 1

    try:
        # Process PDF files first with full parallelism
        process_batch(pdf_files, num_workers, "PDF")
        
        # Process LibreOffice files with limited parallelism
        # LibreOffice has issues with too many concurrent instances
        process_batch(libreoffice_files, libreoffice_workers, "LibreOffice")
        
    finally:
        # Cleanup LibreOffice processes and temp profiles after processing
        kill_libreoffice_processes()
        cleanup_temp_profiles()

    return stats


def main():
    # Register signal handlers and atexit for cleanup
    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)
    atexit.register(kill_libreoffice_processes)
    atexit.register(cleanup_temp_profiles)
    
    parser = argparse.ArgumentParser(
        description="Convert RTF/PDF/DOCX/DOC files to TXT",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument(
        "input_dir",
        type=str,
        help="Directory containing source files",
    )
    parser.add_argument(
        "--output-dir",
        type=str,
        required=True,
        help="Directory for output TXT files",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Overwrite existing TXT files",
    )
    parser.add_argument(
        "--workers",
        type=int,
        default=16,
        help="Number of parallel workers for PDF files (default: 16)",
    )
    parser.add_argument(
        "--libreoffice-workers",
        type=int,
        default=2,
        help="Number of parallel workers for LibreOffice files (default: 2, max recommended: 4)",
    )

    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    output_dir = Path(args.output_dir)

    try:
        soffice_path = get_libreoffice_path()
        print(f"Using LibreOffice at: {soffice_path}")
    except FileNotFoundError as e:
        print(f"Warning: {e}")
        print("RTF/DOCX/DOC conversion will fail; PDF conversion may still work.")
        soffice_path = ""

    try:
        stats = process_directory(
            input_dir,
            output_dir,
            soffice_path,
            skip_existing=not args.overwrite,
            num_workers=args.workers,
            libreoffice_workers=args.libreoffice_workers,
        )

        print("\n" + "=" * 50)
        print("CONVERSION SUMMARY")
        print("=" * 50)
        print(f"Total files found:      {stats['total']}")
        print(f"Successfully converted: {stats['successful']}")
        print(f"Failed conversions:     {stats['failed']}")
        print(f"Skipped (existing):     {stats['skipped']}")
        print("=" * 50)

        exit_code = 1 if stats["failed"] > 0 else 0
        if exit_code == 0:
            print("\nConversion completed!")
    except KeyboardInterrupt:
        print("\n\nInterrupted by user!")
        exit_code = 130
    except Exception as e:
        print(f"\n\nUnexpected error: {e}")
        exit_code = 1
    finally:
        # Always cleanup LibreOffice processes and temp profiles
        print("\nCleaning up...")
        kill_libreoffice_processes()
        cleanup_temp_profiles()
        print("Cleanup completed.")
    
    sys.exit(exit_code)


if __name__ == "__main__":
    main()

